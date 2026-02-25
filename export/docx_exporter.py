from datetime import datetime
from typing import List, Dict
from docx import Document
from docx.shared import Pt

def _add_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)

def _add_bullets(doc: Document, items: List[str]):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(it)

def export_tailored_docx(
    output_path: str,
    base_resume_text: str,
    matched_skills: List[str],
    jd_keywords: List[str],
    suggestions: List[str],
):
    doc = Document()

    _add_heading(doc, "Tailored Additions (generated locally)")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    _add_heading(doc, "Skills matched to JD (from your resume)")
    safe_skills = [s.replace("/", " or ") for s in matched_skills]
    doc.add_paragraph(", ".join(safe_skills) if safe_skills else "None detected")

    _add_heading(doc, "Top JD Keywords (reference only)")
    safe_kw = [k.replace("/", " or ") for k in jd_keywords]
    doc.add_paragraph(", ".join(safe_kw) if safe_kw else "None")

    _add_heading(doc, "Safe suggestions")
    safe_sugg = [s.replace("/", " or ") for s in suggestions]
    _add_bullets(doc, safe_sugg)

    doc.add_page_break()

    _add_heading(doc, "Base Resume (unchanged)")
    # Preserve as plain paragraphs; your resume already has bullets with special chars
    for line in base_resume_text.splitlines():
        doc.add_paragraph(line)

    doc.save(output_path)
